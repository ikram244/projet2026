# generate_docx_chap3.py
from docx import Document
from docx.shared import Inches
import pandas as pd
import os

FIG_DIR = "report_figures"
METRICS_CSV = os.path.join(FIG_DIR, "performance_forward.csv")
OUT_DOCX = "Chapitre_3_PFE.docx"

doc = Document()
doc.add_heading("Chapitre 3 — Étude théorique, état de l'art et environnement de travail", level=1)

# Insert the prepared chapter text (shortened) — replace with full text as needed
doc.add_heading("I. Introduction", level=2)
doc.add_paragraph("Le présent chapitre a pour objectif d’exposer les bases théoriques ... (texte complet à coller)")

# Insert performance table
doc.add_heading("Résultats : performances des modèles forward", level=2)
if os.path.exists(METRICS_CSV):
    perf = pd.read_csv(METRICS_CSV, index_col=0)
    t = doc.add_table(rows=1, cols=len(perf.columns)+1)
    hdr_cells = t.rows[0].cells
    hdr_cells[0].text = "Échelon"
    for i, c in enumerate(perf.columns):
        hdr_cells[i+1].text = c
    for idx, row in perf.iterrows():
        cells = t.add_row().cells
        cells[0].text = str(idx)
        for i, c in enumerate(perf.columns):
            cells[i+1].text = str(round(row[c],3)) if isinstance(row[c], float) else str(int(row[c]))
else:
    doc.add_paragraph("Fichier de métriques introuvable. Exécuter d'abord generate_figures.py.")

# Insert images if present
doc.add_heading("Figures", level=2)
for fname in sorted(os.listdir(FIG_DIR)):
    if fname.endswith(".png"):
        doc.add_paragraph(fname)
        doc.add_picture(os.path.join(FIG_DIR, fname), width=Inches(6))
        doc.add_page_break()

doc.save(OUT_DOCX)
print("Document Word généré:", OUT_DOCX)